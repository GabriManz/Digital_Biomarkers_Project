import os
from docx import Document
from docx.shared import Inches, Pt
from docx.dml.color import RGBColor

def create_summary_docx():
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Estilo de fuentes del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título Principal
    title = doc.add_paragraph()
    run = title.add_run("Resumen de Defensa: Proyecto Digital Biomarkers\n")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(20, 80, 120)
    
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Estructura Conceptual y Guía Sin Datos Específicos (10 Diapositivas)")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph("-" * 80)
    
    slides = [
        {
            "num": "Slide 1: Portada y Objetivo Clínico",
            "speaker": "Integrante 1",
            "content": "Definir el objetivo del proyecto: evaluar la Respuesta Broncodilatadora (BDR) de forma acústica pasiva. Se introduce el biomarcador digital Delta CAS (Tasa CAS Pre-BD menos Tasa CAS Post-BD). El éxito clínico se define como la obtención de deltas positivos para respondedores (BDR+) frente a deltas planos en no-respondedores (BDR-) y controles."
        },
        {
            "num": "Slide 2: Preprocesamiento de Señales y Segmentación",
            "speaker": "Integrante 1",
            "content": "Cadena de filtrado exprés para limpiar ruidos de la señal (remuestreo, filtro Butterworth paso banda y Comb Notch multiarmónico para atenuar zumbidos eléctricos). Se obtienen los segmentos de inspiración y espiración basados en anotaciones clínicas para el modelado."
        },
        {
            "num": "Slide 3: Bloque SOTA (4 Características Físicas)",
            "speaker": "Integrante 2",
            "content": "Explicación de las 4 variables físicas de la literatura acústica (Índice Tonal, Entropía de Picos Espectrales, Curtosis Espectral y la ratio f50/f90). Se justifica por qué un modelo entrenado únicamente con 4 features físicas resulta demasiado limitado e insuficiente a nivel de segmento."
        },
        {
            "num": "Slide 4: Experimento de Deep Learning (CNN en Espectrogramas)",
            "speaker": "Integrante 2",
            "content": "Intento de clasificación de imágenes de espectrogramas Mel mediante redes profundas (CNN y BiLSTM). Explicar que debido al tamaño reducido de la cohorte clínica en validación Leave-One-Subject-Out (LOSO), estas redes complejas sufren un fuerte sobreajuste (overfitting), memorizando el timbre particular del paciente en lugar de la patología."
        },
        {
            "num": "Slide 5: Ingeniería de Características Clásica (137 features)",
            "speaker": "Integrante 3",
            "content": "Justificar la ampliación a un set acústico clásico enriquecido (wavelets, ratios frecuenciales y descriptores temporales). La justificación clínica es analizar la textura acústica completa del audio para capturar transitorios de sibilancia rápidos y sutiles que las variables SOTA no lograban aislar."
        },
        {
            "num": "Slide 6: Resultados y Modelado Clásico",
            "speaker": "Integrante 3",
            "content": "Presentación de resultados del bloque clásico. Explicar que el mayor volumen de descriptores incrementa la exactitud de segmento y que, al promediarlos por sujeto, el biomarcador Delta CAS resultante roza la significación clínica en el test estadístico (p-valor cercano a 0.05), validando el camino tomado."
        },
        {
            "num": "Slide 7: Pipeline Híbrido Optimizado y Bucle LOSO",
            "speaker": "Integrante 4",
            "content": "Fusión del set de características (las clásicas más las SOTA) y la validación cruzada LOSO. Se destaca la honestidad metodológica de recalcular la selección de variables (SelectKBest) dentro de cada fold del training set para evitar cualquier fuga de información clínica."
        },
        {
            "num": "Slide 8: Benchmark de Modelos a Nivel de Segmento",
            "speaker": "Integrante 4",
            "content": "Comparativa de los 9 clasificadores evaluados en el pipeline híbrido. Se justifica la selección de los mejores modelos (Random Forest/Ensemble) por su excelente especificidad a nivel de segmento, prioritaria en clínica para minimizar falsas alarmas acústicas."
        },
        {
            "num": "Slide 9: Biomarcador Clínico Delta CAS y Cancelación de Ruido",
            "speaker": "Integrante 5",
            "content": "Éxito clínico del pipeline híbrido: la prueba estadística es altamente significativa (p-valor < 0.01) y los gráficos (boxplot y scatter plot Pre vs Post) muestran la separación limpia de los grupos respondedores, no-respondedores y controles sanos. Reflexión fundamental: aunque el clasificador segmentario tenga fallos individuales, al promediar sobre cientos de ciclos por paciente (ley de los grandes números), los errores aleatorios se cancelan mutuamente, resultando en un biomarcador Delta CAS a nivel de paciente potente, robusto y clínicamente fiable."
        },
        {
            "num": "Slide 10: Limitaciones del Modelo y Conclusiones",
            "speaker": "Integrante 5",
            "content": "Autocrítica constructiva del estudio: tamaño reducido de la muestra de pacientes clínicos, peligro de sobreajuste de Deep Learning y posible inestabilidad del biomarcador ante tasas basales pequeñas. Como trabajo futuro, se propone aplicar aprendizaje semi-supervisado (Label Propagation) para aprovechar las miles de señales no etiquetadas del dataset."
        }
    ]
    
    for s in slides:
        p_num = doc.add_paragraph()
        run_num = p_num.add_run(s["num"])
        run_num.font.bold = True
        run_num.font.size = Pt(13)
        run_num.font.color.rgb = RGBColor(20, 80, 120)
        
        p_sp = doc.add_paragraph()
        run_sp = p_sp.add_run(f"Ponente: {s['speaker']}")
        run_sp.font.bold = True
        run_sp.font.size = Pt(10.5)
        run_sp.font.color.rgb = RGBColor(120, 30, 30)
        
        doc.add_paragraph(s["content"])
        doc.add_paragraph("-" * 40)
        
    output_path = os.path.join(root_dir, "Resumen_Defensa.docx")
    doc.save(output_path)
    print(f"Documento Word generado exitosamente en: {output_path}")

if __name__ == "__main__":
    root_dir = r"c:\DATA\01_Proyectos\Master\Digital_Biomarkers\Project"
    create_summary_docx()
